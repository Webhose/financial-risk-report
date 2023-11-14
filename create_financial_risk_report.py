import docx
import requests
from openai import OpenAI
import os
import openai
from Levenshtein import ratio
from docx.shared import Pt
from bs4 import BeautifulSoup
import io
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


WEBZ_API_KEY = os.getenv("WEBZ_API_KEY")
openai.api_key = os.getenv("OPENAI_API_KEY")
NUM_OF_REPORTS = 5

client = OpenAI()

def are_similar(str1, str2, threshold=1):
    """
    Check if two strings are similar based on Levenshtein ratio.
    """
    return ratio(str1, str2) > threshold

def remove_similar_strings(articles):
    unique_articles = []
    for article in articles:
        if not any(are_similar(article['text'], existing['text'], 0.7) for existing in unique_articles):
            unique_articles.append(article)
    return unique_articles


def trim_string(string, max_length):

    if len(string) > max_length:
        return string[:max_length]
    else:
        return string

# Function to get news articles from Webz.io API
def fetch_articles(query, api_key, total):
    endpoint = f"https://api.webz.io/filterWebContent?token={api_key}&format=json&q={query}&size=100&ts=0"
    all_posts = []

    while total > 0:
        response = requests.get(endpoint)
        data = response.json()

        posts = data["posts"]
        all_posts.extend(posts)

        total -= len(posts)

        if total > 0 and "next" in data:
            endpoint = f"https://api.webz.io{data['next']}"
        else:
            break

    articles = []
    for article in all_posts:
        article = {'title': article["title"],
                'text': trim_string(trim_title(article["title"]) + "\n\n" + article["text"], 10000),
                'link': article['url'],
                'published': article['published']}

        articles.append(article)

    return articles

def trim_title(input_string):
    words = input_string.split()

    if "|" in input_string:
        return input_string.split("|")[0]

    last_dash_index = input_string.rfind("-")
    if last_dash_index != -1:
        right_of_dash = input_string[last_dash_index + 1:]
        right_words = right_of_dash.split()
        if len(right_words) <= 3 and len(words) > 10:
            return input_string[:last_dash_index]

    return input_string



def add_image_from_base64(doc, image_url):
    response = requests.get(image_url)

    # Check if the request was successful
    if response.status_code == 200:
        image_stream = io.BytesIO(response.content)
        doc.add_picture(image_stream, width=docx.shared.Inches(6))
    else:
        print(f"Failed to download image. Status code: {response.status_code}")


def html_to_word(doc, html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['b', 'ul']):
        if element.name == 'b':
            # Add bold text as a heading
            doc.add_paragraph(element.get_text(), style='Heading 2')
        elif element.name == 'ul':
            for item in element.find_all('li'):
                # Add list items
                doc.add_paragraph(item.get_text(), style='List Bullet')



def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id,)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def insert_titles_in_text(text, reports):
    # Placeholder for inserting the titles
    placeholder = "[]"

    # Extracting the titles from the reports and formatting them with new lines
    titles = "\n".join([report['title'] for report in reports])

    # Replacing the placeholder with the titles
    updated_text = text.replace(placeholder, titles)

    return updated_text

def generate_article_image():
    print("Generating post image")
    image_url = ""
    try:
        response = client.images.generate(
            model="dall-e-3",
            prompt="Create a realistic featured image for a weekly blog post about financial risk reports. The image should depict a modern office environment with a large, clear display screen in the background showing graphs, random companies logos and financial data. In the foreground, arrange a series of  different, but related, documents or tablets, each representing a different financial risk report. These documents should be partially overlapped to convey a sense of abundance and detail. Include elements like pens, glasses, and other office accessories to add to the realism. The overall tone should be professional and sophisticated, with a color scheme that suggests seriousness and reliability, such as shades of blue, grey, and white.",
            n=1,
            size="1024x1024"
        )
        image_url = response.data[0].url



    except Exception as e:
        print("An error occurred generating the image:", str(e))

    return image_url

def get_unique_posts_from_webz(query):
    print("Fetch posts from Webz.io")
    articles = fetch_articles(query, WEBZ_API_KEY, 100)
    filtered_articles = remove_similar_strings(articles)
    return filtered_articles


def call_gpt_completion(prompt):
    return client.chat.completions.create(
        model="gpt-4-1106-preview",
        max_tokens=4096,
        messages=[
            {"role": "user", "content": prompt},
        ]
    )


def generate_reports(filtered_articles):

    print("Generating Reports")

    reports = []

    for article in filtered_articles:

        print(f"Creating report about: {article['title']}")

        prompt = f"""Carefully review the following negative news article in the 'Economy, Business, and Finance' category and determine if there is an explicit financial risk emerging from its content. The article is as follows:

                [ 
                 {article['text']}
                ]
                If the article explicitly mentions or clearly implies a financial risk, generate a detailed financial risk analysis report in HTML format. Use <B> tags to highlight the titles of each section and <UL> and <LI> tags for listing items. The report should include the following sections:

                <HTML>
                <B>1. Executive Summary:</B>

                <UL>
                  <LI>Summarize the main points and the explicit financial risk identified in the article.</LI>
                </UL>
                <B>2. Background Information:</B>

                <UL>
                  <LI>Provide background on the event or issue, focusing on aspects related to the identified financial risk.</LI>
                </UL>
                <B>3. Key Data Extracted:</B>

                <UL>
                  <LI>List key figures, statistics, and significant quotes that are relevant to the financial risk.</LI>
                </UL>
                <B>4. Market and Economic Indicators Impacted:</B>

                <UL>
                  <LI>Discuss the impact on financial markets and economic indicators as it relates to the identified risk.</LI>
                </UL>
                <B>5. Industry-Specific Impact:</B>

                <UL>
                  <LI>Detail the effects on industries, specifically in relation to the financial risk highlighted in the article.</LI>
                </UL>
                <B>6. Company-Specific Impact:</B>

                <UL>
                  <LI>If specific companies are mentioned in the context of the financial risk, explain how the event impacts them.</LI>
                </UL>
                <B>7. Regulatory and Compliance Implications:</B>

                <UL>
                  <LI>Mention any regulatory or compliance issues related to the financial risk.</LI>
                </UL>
                <B>8. Risk Assessment:</B>

                <UL>
                  <LI>Assess the financial risks, focusing on those explicitly mentioned or implied in the article.</LI>
                </UL>
                <B>9. Mitigation Strategies and Recommendations:</B>

                <UL>
                  <LI>Suggest strategies for mitigating the identified financial risks, based on the article.</LI>
                </UL>
                <B>10. Conclusion:</B>

                <UL>
                  <LI>Conclude with the overall implications of the identified financial risk.</LI>
                </UL>
                </HTML>
                If the article does not explicitly mention or imply a financial risk, please respond with: can't produce report.
            """

        try:
            response = call_gpt_completion(prompt)
            report = {'text': ''}

            for choice in response.choices:
                report['text'] += choice.message.content

            if "Executive Summary" in report['text']:
                report['link'] = article['link']
                report['title'] = article['title']
                report['published'] = article['published']
                reports.append(report)
                print(f"Created a report about: {article['title']}")
            else:
                print(f"Can't product report for: {article['title']}")

            if len(reports) == NUM_OF_REPORTS:
                break


        except Exception as e:
            print("An error occurred:", str(e))

    return reports

def generate_intro(reports):
    print("Generate post intro")

    prompt = """
        Write a paragraph introducing a weekly blog post that contains financial risk reports about the following titles, don't elaborate on these titles:
        []

        The reports are created automatically on a weekly basis using Webz.io news api and ChatGPT. The report is generated by calling the Webz.io news API for negative sentiment news articles categorized as "Economy, Business and Finance". The matching news articles are then run through a ChatGPT prompt to analyze if there is a financial risk in the article. If so it create a structured financial risk report. The following weekly post includes up to 5 reports.

        """
    prompt = insert_titles_in_text(prompt, reports)
    intro = ""
    try:
        response = call_gpt_completion(prompt)

        for choice in response.choices:
            intro += choice.message.content
    except Exception as e:
        print("An error occurred:", str(e))

    return intro


def generate_title(intro):
    print("Creating a title")

    prompt = "Create a title using the following text as a context:\n" + intro
    title_text = ""
    try:
        response = call_gpt_completion(prompt)

        for choice in response.choices:
            title_text += choice.message.content
    except Exception as e:
        print("An error occurred:", str(e))

    title_text = title_text.strip('\"')
    if title_text.startswith("Title:"):  # Sometimes ChatGPT prefix the title with Title:
        return title_text[len("Title:"):]

    return title_text


def create_word_doc(file_name, title_text, image_url, intro, reports):

    print("Saving to word document")
    doc = docx.Document()

    # Add a title
    title = doc.add_paragraph()
    title.style = 'Title'
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(24)  # Set the font size
    title_run.font.name = 'NeueHaasUnica-Light'  # Set the font
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the title

    if len(image_url) > 0:
        add_image_from_base64(doc, image_url)
    doc.add_paragraph(intro)

    # Add each report
    for report in reports:
        p = doc.add_paragraph(style='Heading 1')
        add_hyperlink(p, report['link'], report['title'])
        doc.add_paragraph(f"Published on: {report['published']}")
        html_to_word(doc, report['text'])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'NeueHaasUnica-Light'

    # Save the document
    doc.save(file_name)

def main():

    image_url = generate_article_image()
    filtered_articles = get_unique_posts_from_webz("""category:"Economy, Business and Finance" num_chars:>1000 sentiment:negative language:english published:>now-7d  social.facebook.likes:>0""")
    reports = generate_reports(filtered_articles)
    intro = generate_intro(reports)
    title_text = generate_title(intro)
    create_word_doc("financial risk digest.docx", title_text, image_url, intro, reports)


if __name__ == "__main__":
    main()


